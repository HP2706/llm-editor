from typing import Union, Optional, Type, Iterable, AsyncGenerator
from .dataModels import Metadata, Document # type: ignore
from fastapi import UploadFile
import docx2txt
import re
from .utils import count_tokens
from docx import Document as DocxDocument
from io import BytesIO

async def Process_file(file: UploadFile) -> Union[Type[Document], str]:
    '''This function takes a document and returns a list of proposed edits'''

    from docx import Document as DocxDocument
    from io import BytesIO

    if file.filename.endswith('docx'):
        try:
            content_bytes = await file.read()
            file_like_object = BytesIO(content_bytes)
            # Load the .docx file with python-docx
            doc = DocxDocument(file_like_object)
            # Extract text from each paragraph in the document
            full_text = [paragraph.text for paragraph in doc.paragraphs]
            # Join the list of text into a single string
            doc_text = '\n'.join(full_text)
            n_tokens = count_tokens(doc_text)
            n_words = len(doc_text.split())
            metadata = Metadata(title=file.filename, n_tokens=n_tokens, n_words=n_words)
            print("\n\n\n\n\nmetadata", metadata)
            print("doc", doc_text)
            return Document(text=doc_text, metadata=metadata)
        except Exception as e:
            return f"Error processing the file '{file.filename}': {e}"
    elif file.filename.endswith('md') or file.filename.endswith('txt'): # type: ignore
        try:
            content_bytes = await file.read()
            content = content_bytes.decode('utf-8')
            content = ' '.join(content.splitlines())
            pattern = r'\[\[(.*?)\]\]'
            content = re.sub(pattern, r'\1', content)
            n_tokens = count_tokens(content)
            n_words = len(content.split())
            metadata = Metadata(
                title=file.filename if file.filename else "undefined_file_name", 
                n_tokens=n_tokens, n_words=n_words
            )
            return Document(text=content, metadata=metadata)
        except Exception as e:
            return f"Error processing the file '{file.filename}': {e}"
    else:
        return f"Unsupported file format: '{file.filename}'"